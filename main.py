from matplotlib import pyplot as plt
from services import general as g
from src import data_manipulation as dm, ml_models as ml, data_visualisation as dv
from docx import Document
from docx.shared import Cm
import pandas as pd
import numpy as np
from io import BytesIO

# Clean and aggregate data, then snythesise all data into an analytics base table
weather_means = dm.aggregate_gridded_weather_data()
dsa_refined = dm.dsa_data_quality()
form_base = dm.compile_weather_and_energy_data(weather_means, dsa_refined)
abt_hourly = dm.intergrate_daylight_status(form_base)
abt_hourly.to_csv('ABT.csv')


# Group the analytics base table in ways which will allow useful visualisation
abt_hourly.index = pd.to_datetime(abt_hourly.index, errors='coerce')
agg_settings = {'Demand': 'sum', 'Generation': 'sum', 'Import': 'sum', 'Solar': 'sum', 'Wind': 'sum',
                'Other': 'sum', 'Day': 'min', 'Hour': 'min', 'Week': 'min', 't_mean_2m_1h:C': 'mean',
                'wind_speed_mean_10m_1h:ms': 'mean', 'daylight_status': 'min'}
daily_data = abt_hourly.resample('d').agg(agg_settings)
weekly_data = abt_hourly.resample('W').agg(agg_settings)
average_day_hourly = abt_hourly.groupby(['Hour']).agg(agg_settings)
del daily_data['daylight_status'], weekly_data['daylight_status'], \
    average_day_hourly['daylight_status'], average_day_hourly['Week'], average_day_hourly['Day']


# Create descriptive visuals of the data
document = Document()
document.add_heading('Data Visualisations', 0)

# Show energy totals for each day and each week in the dataset
weekly_demand = dv.demand_over_time(weekly_data, 'Weekly demand totals', 'Week number')
document.add_picture(weekly_demand, width=Cm(17), height=Cm(9.5))
daily_demand = dv.demand_over_time(daily_data, 'Daily demand totals', 'Day number')
document.add_picture(daily_demand, width=Cm(17), height=Cm(9.5))

# Show how the variables in the base dataset relate to each other
title = 'Weekly demand totals compared with\nstacked import and generation'
demand_comp = dv.paired_stacked_bars(weekly_data, ['Demand'], ['Import', 'Generation'], title, 'Energy (MW)')
document.add_picture(demand_comp, width=Cm(17), height=Cm(11))
title = 'Weekly generation totals compared with\nstacked solar, wind and other generation'
generat_comp = dv.paired_stacked_bars(weekly_data, ['Generation'], ['Other', 'Wind', 'Solar'], title, 'Generation (MW)')
document.add_picture(generat_comp, width=Cm(17), height=Cm(11))

# Plot array to show how various variables vary through a typical day
fig, axis = plt.subplots(1, 3, figsize=(20, 7), sharex='none', sharey='none')
variables, title, xlabel = ['Demand', 'Generation', 'Import'], 'Energy Demand, Import and Generation', 'Hour of day'
dv.subplot_line_set(variables, axis, 0, title, xlabel, 'Energy(MW)', [0, 24], variables, [0, 25, 2], average_day_hourly)
variables, title, xlabel = ['Solar', 'Wind', 'Other'], 'Generation Components', 'Hour of day'
dv.subplot_line_set(variables, axis, 1, title, xlabel, 'Energy(MW)', [0, 24], variables, [0, 25, 2], average_day_hourly)
variables, title, xlabel = ['t_mean_2m_1h:C', 'wind_speed_mean_10m_1h:ms'], 'Wind and Temperature', 'Hour of day'
dv.subplot_line_set(variables, axis, 2, title, xlabel, 'Temp(C)', [0, 24], variables, [0, 25, 2], average_day_hourly)
memfile = BytesIO()
plt.savefig(memfile)
document.add_picture(memfile, width=Cm(17), height=Cm(4))

# Create standard deviation and r squared sets for various variables for each hour of the day
grouped_hours = [group.drop_duplicates() for _, group in abt_hourly.groupby('Hour')]
sdv_vars = ['Demand', 'Generation', 'Import', 'Solar', 't_mean_2m_1h:C', 'wind_speed_mean_10m_1h:ms', 'daylight_status']
rsq_inds = ['Solar', 'Day', 't_mean_2m_1h:C', 'wind_speed_mean_10m_1h:ms']
std_dev_inds = [[np.std(hour[variable])/np.mean(hour[variable]) for hour in grouped_hours] for variable in sdv_vars]
rsq_inds_vs_demand, rsq_inds_vs_import = [], []
for ind_var in rsq_inds:
    rsqs_demand = [g.coeff_of_determination(hour[ind_var], hour['Demand']) for hour in grouped_hours]
    rsqs_import = [g.coeff_of_determination(hour[ind_var], hour['Import']) for hour in grouped_hours]
    rsq_inds_vs_demand.append(rsqs_demand)
    rsq_inds_vs_import.append(rsqs_import)

# Show the above
fig2, axis = plt.subplots(1, 3, figsize=(20, 7), sharex='none', sharey='none')
dv.subplot_line_set([_ for _ in std_dev_inds], axis, 0, 'Standard Deviation for Various Variables Throughout the Day',
                    'Hour of day', 'Coefficient of Standard Deviation', [0, 24], sdv_vars, [0, 25, 2])
dv.subplot_line_set([_ for _ in rsq_inds_vs_demand], axis, 1, 'R Squared for Various Variables Throughout the Day',
                    'Hour of day', 'R Squared', [0, 24], rsq_inds, [0, 25, 2])
dv.subplot_line_set([_ for _ in rsq_inds_vs_import], axis, 2, 'R Squared for Various Variables Throughout the Day',
                    'Hour of day', 'R Squared', [0, 24], rsq_inds, [0, 25, 2])
memfile = BytesIO()
plt.savefig(memfile)
document.add_picture(memfile, width=Cm(17), height=Cm(4))


# Generate models, deploy on test data and check accuracy
document.add_heading('ML Model Test Results', 0)
path = 'ABT.csv'

# Define training/test threshold and sum the demand and import levels that occured
training_threshold = 2498
actual_demand = sum(list(abt_hourly['Demand'])[training_threshold:])
actual_import = sum(list(abt_hourly['Import'])[training_threshold:])

# Create models and generate predictions
mr_demand = ml.multireg(path, ['t_mean_2m_1h:C', 'Hour', 'Day'], 'Demand', training_threshold)
mr_import = ml.multireg(path, ['t_mean_2m_1h:C', 'Hour', 'Day'], 'Import', training_threshold)
ann_demand = ml.ann(path, [1, 7, 8, 10], [1, 2, 3], 0, training_threshold)
ann_import = ml.ann(path, [3, 7, 8, 10], [1, 2, 3], 0, training_threshold)

# Write results and accuracy metrics, for each model and dependent variable, to word document
skl_title = 'SkLearn multiple regression'
document.add_paragraph(g.ml_performance_output(skl_title, 'Demand', actual_demand, mr_demand), style='List Number')
document.add_paragraph(g.ml_performance_output(skl_title, 'Import', actual_import, mr_import), style='List Number')
document.add_paragraph(g.ml_performance_output('Keras ANN', 'Import', actual_import, ann_import), style='List Number')
document.add_paragraph(g.ml_performance_output('Keras ANN', 'Demand', actual_demand, ann_demand), style='List Number')
document.save('Data Visualisation and ML Model Testing.docx')
